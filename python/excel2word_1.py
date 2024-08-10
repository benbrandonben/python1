import pandas as pd
from docx import Document
from openpyxl import load_workbook
 
# 加载Excel文件
workbook = load_workbook(r"D:\Program Files\Python310\HCT_Astra6_BSW系统需求表 Rp09.xlsx")
sheets = workbook.sheetnames
 
# 读取Excel文件
#df = pd.read_excel('HCT_Astra_J6E_HRBT_需求-V1.00.xlsx', sheet_name='H01_最小系统')

sheet_num = len(sheets)
#print(sheet_num)
#sheet = workbook.worksheets[0]
# 创建一个Word文档
doc = Document()
# 遍历DataFrame的行
#for sheet in workbook:
for m in range (2,sheet_num):
    #将MCU_SOC协议的sheet放到最后，该sheet没有内容
    if m<=sheet_num-2:
        row_count = workbook.worksheets[m].max_row
        print(row_count)
        column_count = workbook.worksheets[m].max_column
        #从sheet的第1行开始遍历直到最后一行
        for i in range (1, row_count):
            #for j in range (1 , column_count):
            cell_value = workbook.worksheets[m].cell(row=i, column= 1).value
            #第1行第一列为2时，将B列内容作为sheet标题写入word
            if cell_value ==2 and i == 1:
            #if cell_value ==2:
                #print('111')
                write_value = workbook.worksheets[m].cell(row=i, column= 2).value
                doc.add_paragraph(write_value)
                i=i+3
            #获取A列的值，为1时记录该行B列的内容并加入到word文档中
            elif cell_value ==1:
                write_value = workbook.worksheets[m].cell(row=i, column= 2).value
                doc.add_paragraph(write_value)
                i=i+1
            #A列的值是3时，将D列的需求内容记录到word中
            elif cell_value ==3:
                write_value2 = workbook.worksheets[m].cell(row=i, column= 4).value
                doc.add_paragraph(f'需求内容: {write_value2}')
                i=i+1
            #A列的值不为1时则为标题下的逐条内容，需要将需求标题及需求内容记录到word中
            else :
                write_value1 = workbook.worksheets[m].cell(row=i, column= 3).value
                write_value2 = workbook.worksheets[m].cell(row=i, column= 4).value
                doc.add_paragraph(f'需求标题: {write_value1}')
                doc.add_paragraph(f'需求内容: {write_value2}')
    else:
        write_value = workbook.worksheets[m].cell(row=i, column= 2).value
        doc.add_paragraph(write_value)

# 保存Word文档
doc.save('output.docx')



'''
    # 添加标题到Word文档
    doc.add_heading(f'{name}\'s Information', 0)
    
    # 添加文本内容到Word文档
    doc.add_paragraph(f'Name: {name}')
    doc.add_paragraph(f'Age: {age}')
    doc.add_paragraph(f'Address: {address}')
    
    # 在每个人的信息后添加一个分隔符
    doc.add_paragraph('--------------------')
 
# 保存Word文档
doc.save('output.docx')
'''